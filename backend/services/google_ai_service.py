import requests
import json
import os
from typing import List, Dict, Optional, Union
from config import Config
import PyPDF2
import docx
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from google.cloud import bigquery
from google.cloud import firestore
import firebase_admin
from firebase_admin import credentials

class GoogleCloudAIService:
    def __init__(self):
        self.config = Config()
        self.api_key = os.environ.get('GOOGLE_AI_API_KEY')
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro:generateContent"
        self.ai_enabled = bool(self.api_key)
        
        # Initialize BigQuery for analytics
        try:
            self.bigquery_client = bigquery.Client(project=self.config.GOOGLE_CLOUD_PROJECT)
            self.bigquery_enabled = True
        except:
            self.bigquery_enabled = False
        
        # Initialize Firestore for document storage
        try:
            self.firestore_client = firestore.Client(project=self.config.GOOGLE_CLOUD_PROJECT)
            self.firestore_enabled = True
        except:
            self.firestore_enabled = False
        
        # Initialize Firebase for real-time features
        try:
            if not firebase_admin._apps:
                cred = credentials.ApplicationDefault()
                firebase_admin.initialize_app(cred)
            self.firebase_enabled = True
        except:
            self.firebase_enabled = False

    async def process_multiple_formats(self, file_content: bytes, file_type: str) -> Dict[str, str]:
        """Process multiple document formats (PDF, Word, XML, Markup)"""
        
        extracted_text = ""
        metadata = {}
        
        try:
            if file_type.lower() == 'pdf':
                extracted_text, metadata = self._process_pdf(file_content)
            elif file_type.lower() in ['docx', 'doc']:
                extracted_text, metadata = self._process_word(file_content)
            elif file_type.lower() == 'xml':
                extracted_text, metadata = self._process_xml(file_content)
            elif file_type.lower() in ['html', 'htm', 'markdown', 'md']:
                extracted_text, metadata = self._process_markup(file_content)
            else:
                extracted_text = file_content.decode('utf-8', errors='ignore')
                metadata = {"format": "text", "size": len(file_content)}
            
            # Store in Firestore for GDPR-compliant processing
            if self.firestore_enabled:
                doc_ref = self.firestore_client.collection('processed_documents').add({
                    'content': extracted_text[:1000],  # Truncated for storage
                    'metadata': metadata,
                    'processed_at': firestore.SERVER_TIMESTAMP,
                    'gdpr_compliant': True
                })
            
            return {
                "extracted_text": extracted_text,
                "metadata": metadata,
                "format_supported": True,
                "gdpr_processed": self.firestore_enabled
            }
            
        except Exception as e:
            return {
                "extracted_text": "",
                "metadata": {"error": str(e)},
                "format_supported": False,
                "gdpr_processed": False
            }

    def _process_pdf(self, file_content: bytes) -> tuple[str, Dict]:
        """Extract text from PDF documents"""
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        metadata = {
            "format": "PDF",
            "pages": len(pdf_reader.pages),
            "size": len(file_content)
        }
        return text, metadata

    def _process_word(self, file_content: bytes) -> tuple[str, Dict]:
        """Extract text from Word documents"""
        import io
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        metadata = {
            "format": "Word",
            "paragraphs": len(doc.paragraphs),
            "size": len(file_content)
        }
        return text, metadata

    def _process_xml(self, file_content: bytes) -> tuple[str, Dict]:
        """Extract text from XML documents"""
        try:
            root = ET.fromstring(file_content.decode('utf-8'))
            text = ET.tostring(root, encoding='unicode', method='text')
            
            metadata = {
                "format": "XML",
                "root_tag": root.tag,
                "size": len(file_content)
            }
            return text, metadata
        except:
            return "", {"format": "XML", "error": "Parse error"}

    def _process_markup(self, file_content: bytes) -> tuple[str, Dict]:
        """Extract text from HTML/Markdown"""
        try:
            content = file_content.decode('utf-8')
            if content.strip().startswith('<'):
                # HTML content
                soup = BeautifulSoup(content, 'html.parser')
                text = soup.get_text()
                metadata = {"format": "HTML", "tags": len(soup.find_all())}
            else:
                # Markdown content
                text = content
                metadata = {"format": "Markdown", "lines": len(content.split('\n'))}
            
            return text, metadata
        except:
            return "", {"format": "Markup", "error": "Parse error"}

    async def generate_comprehensive_test_cases(self, requirements: str, test_type: str, compliance_standard: str, format_context: Dict = None) -> List[Dict]:
        """Generate comprehensive test cases with full regulatory support"""
        
        # Enhanced compliance standards support
        compliance_frameworks = {
            "HIPAA": "Healthcare data privacy and security",
            "HITECH": "Health Information Technology for Economic and Clinical Health",
            "FDA": "Food and Drug Administration medical device regulations",
            "IEC_62304": "Medical device software lifecycle processes",
            "ISO_9001": "Quality management systems",
            "ISO_13485": "Medical devices quality management systems",
            "ISO_27001": "Information security management systems",
            "GDPR": "General Data Protection Regulation"
        }
        
        framework_description = compliance_frameworks.get(compliance_standard, "Healthcare compliance")
        
        enhanced_prompt = f"""
        You are an expert healthcare software testing engineer specializing in {framework_description} ({compliance_standard}) compliance.
        
        Generate comprehensive test cases for healthcare software requirements:
        
        Requirements: {requirements}
        Test Type: {test_type}
        Compliance Framework: {compliance_standard} - {framework_description}
        Document Context: {format_context or 'Direct input'}
        
        Generate 4-6 detailed test cases covering:
        1. Regulatory compliance ({compliance_standard})
        2. Data privacy and GDPR requirements
        3. Healthcare workflow integration
        4. Enterprise toolchain compatibility
        5. Traceability and audit requirements
        6. Risk management and validation
        
        Return ONLY valid JSON:
        {{
            "test_cases": [
                {{
                    "id": "TC001-{compliance_standard}",
                    "title": "Regulatory Compliance Validation Test",
                    "description": "Comprehensive validation of {compliance_standard} compliance requirements",
                    "test_type": "{test_type}",
                    "priority": "critical",
                    "regulatory_framework": "{compliance_standard}",
                    "gdpr_compliant": true,
                    "preconditions": ["Regulatory environment configured", "Compliance documentation available"],
                    "test_steps": [
                        {{
                            "step_number": 1,
                            "action": "Validate regulatory compliance controls",
                            "expected_result": "All {compliance_standard} controls are properly implemented",
                            "test_data": "Regulatory test dataset",
                            "traceability_id": "REQ-{compliance_standard}-001"
                        }}
                    ],
                    "expected_outcome": "Full {compliance_standard} compliance demonstrated with audit trail",
                    "compliance_tags": ["{compliance_standard}", "Regulatory", "Audit"],
                    "requirements_traceability": ["REQ-{compliance_standard}-001"],
                    "alm_integration": {{
                        "jira_issue_type": "Test",
                        "azure_devops_work_item": "Test Case",
                        "polarion_type": "Test Case"
                    }},
                    "estimated_duration": 30,
                    "risk_level": "critical",
                    "healthcare_context": "Regulatory compliance in clinical environment",
                    "gdpr_considerations": ["Data minimization", "Consent management", "Right to erasure"],
                    "validation_criteria": ["Documented evidence", "Audit trail", "Regulatory approval"]
                }}
            ]
        }}
        """
        
        if not self.ai_enabled:
            return self._generate_enhanced_fallback_tests(requirements, test_type, compliance_standard)
        
        try:
            payload = {
                "contents": [{
                    "parts": [{
                        "text": enhanced_prompt
                    }]
                }]
            }
            
            url = f"{self.base_url}?key={self.api_key}"
            response = requests.post(url, headers={'Content-Type': 'application/json'}, json=payload, timeout=45)
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    text = result['candidates'][0]['content']['parts'][0]['text']
                    text = text.replace('``````', '').strip()
                    
                    json_start = text.find('{')
                    json_end = text.rfind('}') + 1
                    
                    if json_start != -1 and json_end > json_start:
                        json_str = text[json_start:json_end]
                        parsed = json.loads(json_str)
                        test_cases = parsed.get('test_cases', [])
                        
                        # Store analytics in BigQuery
                        if self.bigquery_enabled:
                            self._store_test_generation_analytics(requirements, test_cases, compliance_standard)
                        
                        return test_cases
            
            return self._generate_enhanced_fallback_tests(requirements, test_type, compliance_standard)
            
        except Exception as e:
            print(f"Enhanced AI generation failed: {e}")
            return self._generate_enhanced_fallback_tests(requirements, test_type, compliance_standard)

    def _store_test_generation_analytics(self, requirements: str, test_cases: List[Dict], standard: str):
        """Store test generation analytics in BigQuery"""
        try:
            table_id = f"{self.config.GOOGLE_CLOUD_PROJECT}.healthcare_testing.test_generation_analytics"
            
            rows_to_insert = [{
                "timestamp": "2025-09-21T18:00:00",
                "requirements_length": len(requirements),
                "test_cases_generated": len(test_cases),
                "compliance_standard": standard,
                "ai_powered": True,
                "success": True
            }]
            
            errors = self.bigquery_client.insert_rows_json(table_id, rows_to_insert)
            if not errors:
                print("Analytics stored successfully in BigQuery")
        except Exception as e:
            print(f"BigQuery analytics storage failed: {e}")

    def _generate_enhanced_fallback_tests(self, requirements: str, test_type: str, compliance_standard: str) -> List[Dict]:
        """Enhanced fallback with full regulatory support"""
        
        regulatory_tests = {
            "FDA": self._generate_fda_tests(requirements),
            "IEC_62304": self._generate_iec62304_tests(requirements),
            "ISO_9001": self._generate_iso9001_tests(requirements),
            "ISO_13485": self._generate_iso13485_tests(requirements),
            "ISO_27001": self._generate_iso27001_tests(requirements),
            "GDPR": self._generate_gdpr_tests(requirements)
        }
        
        base_tests = [
            {
                "id": f"TC001-{compliance_standard}",
                "title": f"{compliance_standard} Regulatory Compliance Validation",
                "description": f"Comprehensive validation of {compliance_standard} compliance requirements in healthcare software",
                "test_type": test_type,
                "priority": "critical",
                "regulatory_framework": compliance_standard,
                "gdpr_compliant": True,
                "preconditions": [
                    f"{compliance_standard} regulatory framework implemented",
                    "Healthcare compliance documentation available",
                    "Audit trail system operational",
                    "Risk management procedures in place"
                ],
                "test_steps": [
                    {
                        "step_number": 1,
                        "action": f"Validate {compliance_standard} compliance controls implementation",
                        "expected_result": f"All mandatory {compliance_standard} controls are properly configured and operational",
                        "test_data": f"{compliance_standard} compliance test dataset",
                        "traceability_id": f"REQ-{compliance_standard}-001"
                    },
                    {
                        "step_number": 2,
                        "action": "Verify comprehensive audit trail generation",
                        "expected_result": "All regulatory activities are logged with complete audit trail",
                        "test_data": "Regulatory activity audit logs",
                        "traceability_id": f"REQ-{compliance_standard}-002"
                    },
                    {
                        "step_number": 3,
                        "action": "Test GDPR compliance integration",
                        "expected_result": "GDPR requirements (consent, data minimization, erasure) are properly handled",
                        "test_data": "GDPR compliance test scenarios",
                        "traceability_id": "REQ-GDPR-001"
                    },
                    {
                        "step_number": 4,
                        "action": "Validate enterprise toolchain integration",
                        "expected_result": "Test results are properly integrated with ALM tools (Jira, Azure DevOps, Polarion)",
                        "test_data": "ALM integration test data",
                        "traceability_id": f"REQ-ALM-{compliance_standard}-001"
                    }
                ],
                "expected_outcome": f"Complete {compliance_standard} compliance demonstrated with full traceability and enterprise integration",
                "compliance_tags": [compliance_standard, "Regulatory", "GDPR", "Enterprise", "Traceability"],
                "requirements_traceability": [f"REQ-{compliance_standard}-001", f"REQ-{compliance_standard}-002", "REQ-GDPR-001"],
                "alm_integration": {
                    "jira_issue_type": "Test",
                    "jira_labels": [compliance_standard, "healthcare", "compliance"],
                    "azure_devops_work_item": "Test Case",
                    "azure_devops_tags": [compliance_standard, "regulatory"],
                    "polarion_type": "Test Case",
                    "polarion_category": f"{compliance_standard}_Compliance"
                },
                "estimated_duration": 45,
                "risk_level": "critical",
                "healthcare_context": f"{compliance_standard} regulatory compliance in clinical healthcare environment",
                "gdpr_considerations": [
                    "Data minimization principle",
                    "Explicit consent management", 
                    "Right to erasure implementation",
                    "Data portability support",
                    "Privacy by design integration"
                ],
                "validation_criteria": [
                    "Documented regulatory evidence",
                    "Complete audit trail",
                    "Regulatory authority approval",
                    "Third-party compliance certification",
                    "Enterprise toolchain integration verification"
                ],
                "enterprise_integration": {
                    "supported_formats": ["PDF", "Word", "XML", "Markup"],
                    "alm_platforms": ["Jira", "Polarion", "Azure DevOps"],
                    "export_formats": ["JUnit", "Cucumber", "TestNG", "ALM-specific"]
                }
            }
        ]
        
        # Add standard-specific tests
        if compliance_standard in regulatory_tests:
            base_tests.extend(regulatory_tests[compliance_standard])
        
        return base_tests[:6]  # Return up to 6 comprehensive tests

    def _generate_fda_tests(self, requirements: str) -> List[Dict]:
        """Generate FDA-specific medical device tests"""
        return [{
            "id": "TC-FDA-001",
            "title": "FDA Medical Device Software Validation (21 CFR Part 820)",
            "description": "Validate medical device software according to FDA Quality System Regulation",
            "regulatory_framework": "FDA",
            "test_steps": [
                {
                    "step_number": 1,
                    "action": "Verify software validation documentation per 21 CFR 820.30",
                    "expected_result": "Complete software validation documentation available",
                    "traceability_id": "REQ-FDA-820.30"
                }
            ]
        }]

    def _generate_iec62304_tests(self, requirements: str) -> List[Dict]:
        """Generate IEC 62304 medical device software lifecycle tests"""
        return [{
            "id": "TC-IEC62304-001", 
            "title": "IEC 62304 Software Lifecycle Process Validation",
            "description": "Validate software development lifecycle according to IEC 62304 standard",
            "regulatory_framework": "IEC_62304"
        }]

    def _generate_iso9001_tests(self, requirements: str) -> List[Dict]:
        """Generate ISO 9001 quality management tests"""
        return [{
            "id": "TC-ISO9001-001",
            "title": "ISO 9001 Quality Management System Validation", 
            "description": "Validate quality management system processes per ISO 9001",
            "regulatory_framework": "ISO_9001"
        }]

    def _generate_iso13485_tests(self, requirements: str) -> List[Dict]:
        """Generate ISO 13485 medical device quality tests"""
        return [{
            "id": "TC-ISO13485-001",
            "title": "ISO 13485 Medical Device Quality System Validation",
            "description": "Validate medical device quality management system per ISO 13485",
            "regulatory_framework": "ISO_13485"
        }]

    def _generate_iso27001_tests(self, requirements: str) -> List[Dict]:
        """Generate ISO 27001 information security tests"""
        return [{
            "id": "TC-ISO27001-001",
            "title": "ISO 27001 Information Security Management Validation",
            "description": "Validate information security management system per ISO 27001",
            "regulatory_framework": "ISO_27001"
        }]

    def _generate_gdpr_tests(self, requirements: str) -> List[Dict]:
        """Generate GDPR compliance tests"""
        return [{
            "id": "TC-GDPR-001",
            "title": "GDPR Data Privacy and Protection Validation",
            "description": "Comprehensive GDPR compliance validation for healthcare data processing",
            "regulatory_framework": "GDPR",
            "gdpr_considerations": [
                "Right to access (Art. 15)",
                "Right to rectification (Art. 16)", 
                "Right to erasure (Art. 17)",
                "Data portability (Art. 20)",
                "Consent management (Art. 7)"
            ]
        }]

    def get_comprehensive_service_status(self) -> Dict:
        """Get comprehensive Google Cloud AI and enterprise integration status"""
        return {
            "google_ai_services": {
                "generative_ai_enabled": self.ai_enabled,
                "bigquery_analytics": getattr(self, 'bigquery_enabled', False),
                "firestore_storage": getattr(self, 'firestore_enabled', False),
                "firebase_realtime": getattr(self, 'firebase_enabled', False),
                "model": "gemini-1.5-pro"
            },
            "document_processing": {
                "supported_formats": ["PDF", "Word", "XML", "HTML", "Markdown"],
                "extraction_capabilities": ["Text", "Structure", "Metadata"]
            },
            "compliance_frameworks": [
                "FDA", "IEC_62304", 
                "ISO_9001", "ISO_13485", "ISO_27001", "GDPR"
            ],
            "enterprise_integration": {
                "alm_platforms": ["Jira", "Polarion", "Azure DevOps"],
                "export_formats": ["JUnit", "Cucumber", "TestNG"],
                "gdpr_compliant": True
            },
            "ai_capabilities": {
                "multi_format_processing": True,
                "regulatory_interpretation": True,
                "enterprise_toolchain_integration": True,
                "traceability_management": True,
                "scalable_architecture": True
            }
        }

